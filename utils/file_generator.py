import os
import subprocess
from pathlib import Path
from typing import Dict

from docx import Document
from docx.shared import Pt

from jinja2 import Environment, FileSystemLoader

from .template_manager import (
    get_template_latex,
    format_contact_info,
    format_experience_latex,
    format_education_latex,
    format_skills_latex,
    format_projects_latex,
    latex_escape,
)


OUTPUT_DIR = Path("outputs")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def generate_docx(data: Dict, template: str) -> str:
    doc = Document()

    # Template-specific style settings
    tpl = (template or "professional").lower()
    if tpl == "modern":
        name_size = 22; heading_size = 14; normal_size = 10
        heading_upper = True; heading_color = (0x15, 0x63, 0xC1)  # blue
    elif tpl == "minimal":
        name_size = 18; heading_size = 11; normal_size = 9
        heading_upper = False; heading_color = (0x44, 0x44, 0x44)  # dark gray
    else:  # professional
        name_size = 20; heading_size = 12; normal_size = 11
        heading_upper = False; heading_color = (0x00, 0x00, 0x00)

    def set_paragraph_font(p, size):
        for run in p.runs:
            run.font.size = Pt(size)

    def add_section_heading(text):
        title = text.upper() if heading_upper else text
        h = doc.add_heading(title, level=1)
        for run in h.runs:
            run.font.size = Pt(heading_size)
            run.font.color.rgb = __import__('docx.shared', fromlist=['RGBColor']).shared.RGBColor(*heading_color)  # type: ignore
        # decorative separator for Modern
        if tpl == "modern":
            sep = doc.add_paragraph()
            sep.add_run("\u2500" * 40).font.size = Pt(6)
        return h

    # Name
    name = (data.get("personal_info", {}).get("name") or "Your Name")
    t0 = doc.add_heading(level=0)
    run = t0.add_run(name)
    run.bold = True
    run.font.size = Pt(name_size)

    # Contact
    contact = format_contact_info(data.get("personal_info", {}))
    if contact:
        p = doc.add_paragraph(contact)
        set_paragraph_font(p, normal_size)

    # Summary
    summary = (data.get("personal_info", {}).get("summary") or "").strip()
    if summary:
        add_section_heading("Summary")
        p = doc.add_paragraph(summary)
        set_paragraph_font(p, normal_size)

    # Skills
    skills = data.get("skills", []) or []
    if skills:
        add_section_heading("Skills")
        if tpl == "minimal":
            p = doc.add_paragraph(" ".join([f"[{s}]" for s in skills]))
        else:
            p = doc.add_paragraph(", ".join(skills))
        set_paragraph_font(p, normal_size)

    # Education
    education = data.get("education", []) or []
    if education:
        add_section_heading("Education")
        for e in education:
            para = doc.add_paragraph(str(e), style="List Bullet" if tpl != "minimal" else None)
            set_paragraph_font(para, normal_size)

    # Experience
    experience = data.get("experience", []) or []
    if experience:
        add_section_heading("Experience")
        for job in experience:
            jt = job.get("title", ""); co = job.get("company", ""); dt = job.get("dates", "")
            header = f"{jt} — {co} ({dt})".strip()
            h = doc.add_paragraph(header)
            set_paragraph_font(h, normal_size)
            for b in job.get("responsibilities", []) or []:
                para = doc.add_paragraph(b, style="List Bullet" if tpl != "minimal" else None)
                set_paragraph_font(para, normal_size)

    # Projects
    projects = data.get("projects", []) or []
    if projects:
        add_section_heading("Projects")
        for p in projects:
            name = p.get("name", "Project"); desc = p.get("description", ""); tech = p.get("technologies", "")
            line = name + (f" ({tech})" if tech else "")
            lp = doc.add_paragraph(line)
            set_paragraph_font(lp, normal_size)
            if desc:
                dp = doc.add_paragraph(desc, style="List Bullet" if tpl != "minimal" else None)
                set_paragraph_font(dp, normal_size)

    out_path = OUTPUT_DIR / "resume.docx"
    doc.save(str(out_path))
    return str(out_path)


def generate_pdf(data: Dict, template: str) -> str:
    try:
        return generate_pdf_latex(data, template)
    except Exception:
        try:
            return generate_pdf_html(data, template)
        except Exception:
            return generate_pdf_from_docx(data, template)


def generate_pdf_latex(data: Dict, template: str) -> str:
    tex_template_path = Path(get_template_latex(template))
    if not tex_template_path.exists():
        raise FileNotFoundError(f"LaTeX template not found: {tex_template_path}")

    env = Environment(loader=FileSystemLoader(str(tex_template_path.parent)))
    env.filters["lex"] = latex_escape
    j2 = env.get_template(tex_template_path.name)

    personal = data.get("personal_info", {})
    rendered = j2.render(
        NAME=(personal.get("name") or "Your Name"),
        CONTACT=format_contact_info(personal),
        SUMMARY=(personal.get("summary") or ""),
        EXPERIENCE=format_experience_latex(data.get("experience", []) or []),
        EDUCATION=format_education_latex(data.get("education", []) or []),
        SKILLS=format_skills_latex(data.get("skills", []) or []),
        PROJECTS=format_projects_latex(data.get("projects", []) or []),
    )

    out_tex = OUTPUT_DIR / "resume.tex"
    out_pdf = OUTPUT_DIR / "resume.pdf"
    with out_tex.open("w", encoding="utf-8") as f:
        f.write(rendered)

    # Compile with pdflatex if available
    if not _is_command_available("pdflatex"):
        raise RuntimeError("pdflatex not available")

    try:
        subprocess.run(
            [
                "pdflatex",
                "-interaction=nonstopmode",
                f"-output-directory={str(OUTPUT_DIR)}",
                str(out_tex),
            ],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
    finally:
        cleanup_temp_files()

    return str(out_pdf)


def generate_pdf_from_docx(data: Dict, template: str) -> str:
    docx_path = generate_docx(data, template)
    pdf_path = OUTPUT_DIR / "resume.pdf"
    try:
        from docx2pdf import convert  # type: ignore

        convert(docx_path, str(pdf_path))
        return str(pdf_path)
    except Exception:
        raise RuntimeError("PDF generation failed (no LaTeX/HTML engine or Word automation available). Download DOCX instead.")


def cleanup_temp_files():
    for ext in (".aux", ".log", ".out", ".toc"):
        for p in OUTPUT_DIR.glob(f"*{ext}"):
            try:
                p.unlink()
            except Exception:
                pass


def _is_command_available(cmd: str) -> bool:
    from shutil import which

    return which(cmd) is not None


def generate_pdf_html(data: Dict, template: str) -> str:
    out_pdf = OUTPUT_DIR / "resume.pdf"
    html = _render_html(data, template)
    try:
        from xhtml2pdf import pisa  # type: ignore
        from io import BytesIO

        result = BytesIO()
        pisa_status = pisa.CreatePDF(html, dest=result)
        if pisa_status.err:
            raise RuntimeError("xhtml2pdf failed to render HTML")
        with open(out_pdf, "wb") as f:
            f.write(result.getvalue())
        return str(out_pdf)
    except Exception as e:
        raise e


def _render_html(data: Dict, template: str) -> str:
    personal = data.get("personal_info", {}) or {}
    name = personal.get("name") or "Your Name"
    contact_parts = []
    for k in ["email", "phone", "linkedin", "github", "website", "location"]:
        v = personal.get(k)
        if v:
            contact_parts.append(str(v))
    contact = " | ".join(contact_parts)
    skills = data.get("skills", []) or []
    education = data.get("education", []) or []
    experience = data.get("experience", []) or []
    projects = data.get("projects", []) or []

    # CSS variants per template
    tpl = (template or "professional").lower()
    if tpl == "modern":
        css = """
    body { font-family: Helvetica, Arial, sans-serif; font-size: 12pt; color: #222; }
    h1 { font-size: 24pt; margin: 0; color: #1563C1; }
    h2 { font-size: 14pt; margin-top: 14pt; border-bottom: 2px solid #1563C1; text-transform: uppercase; }
    .contact { color: #444; margin-top: 2pt; }
    .chips span { display: inline-block; margin: 2px 6px 2px 0; padding: 2px 8px; border: 1px solid #1563C1; border-radius: 12px; font-size: 10pt; }
    .job { margin-bottom: 8pt; }
    .job h3 { margin: 0; font-size: 12pt; }
    .meta { color: #666; font-size: 10pt; }
        """
    elif tpl == "minimal":
        css = """
    body { font-family: Arial, sans-serif; font-size: 11pt; color: #222; }
    h1 { font-size: 18pt; margin: 0; }
    h2 { font-size: 12pt; margin-top: 10pt; color: #444; }
    .contact { color: #555; margin-top: 2pt; }
    .chips span { display: inline-block; margin: 1px 4px 1px 0; padding: 1px 6px; border: 1px solid #bbb; border-radius: 10px; font-size: 9pt; }
    .job { margin-bottom: 6pt; }
    .job h3 { margin: 0; font-size: 11pt; }
    .meta { color: #666; font-size: 9pt; }
        """
    else:  # professional
        css = """
    body { font-family: Georgia, 'Times New Roman', serif; font-size: 12pt; color: #111; }
    h1 { font-size: 22pt; margin: 0; }
    h2 { font-size: 13pt; margin-top: 12pt; border-bottom: 1px solid #ddd; }
    .contact { color: #555; margin-top: 2pt; }
    .chips span { display: inline-block; margin: 2px 6px 2px 0; padding: 2px 8px; border: 1px solid #ccc; border-radius: 12px; font-size: 10pt; }
    .job { margin-bottom: 8pt; }
    .job h3 { margin: 0; font-size: 12pt; }
    .meta { color: #666; font-size: 10pt; }
        """

    html = f"""
<!DOCTYPE html>
<html>
<head>
  <meta charset='utf-8'/>
  <style>
    {css}
    ul {{ margin-top: 4pt; }}
  </style>
  <title>Resume</title>
  </head>
  <body>
    <h1>{latex_escape(name)}</h1>
    <div class='contact'>{latex_escape(contact)}</div>
    {f"<h2>Summary</h2><p>{latex_escape(personal.get('summary') or '')}</p>" if (personal.get('summary') or '') else ''}
    {f"<h2>Skills</h2><div class='chips'>" + ''.join(f"<span>{latex_escape(s)}</span>" for s in skills) + "</div>" if skills else ''}
    {"<h2>Education</h2><ul>" + ''.join(f"<li>{latex_escape(str(e))}</li>" for e in education) + "</ul>" if education else ''}
    {"<h2>Experience</h2>" if experience else ''}
    {''.join(f"<div class='job'><h3>{latex_escape(j.get('title',''))} — {latex_escape(j.get('company',''))}</h3><div class='meta'>{latex_escape(j.get('dates',''))}</div><ul>" + ''.join(f"<li>{latex_escape(b)}</li>" for b in (j.get('responsibilities',[]) or [])) + "</ul></div>" for j in experience)}
    {"<h2>Projects</h2><ul>" + ''.join(f"<li><strong>{latex_escape(p.get('name','Project'))}</strong>: {latex_escape(p.get('description',''))}</li>" for p in projects) + "</ul>" if projects else ''}
  </body>
</html>
"""
    return html
