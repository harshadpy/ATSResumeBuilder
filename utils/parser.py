import re
import PyPDF2
import docx
from io import BytesIO
try:
    import fitz  # type: ignore
except Exception:
    fitz = None  # type: ignore
try:
    from PIL import Image  # type: ignore
except Exception:
    Image = None  # type: ignore
try:
    import pytesseract  # type: ignore
except Exception:
    pytesseract = None  # type: ignore
try:
    import spacy  # type: ignore
except Exception:
    spacy = None  # type: ignore

def extract_text_from_pdf(file):
    """Extract text from PDF file"""
    try:
        data = file.read()
        pdf_reader = PyPDF2.PdfReader(BytesIO(data))
        parts = []
        for page in pdf_reader.pages:
            t = page.extract_text() or ""
            parts.append(t)
        text = "\n".join(parts)
        if len(text.strip()) >= 50 or not fitz or not pytesseract or not Image:
            return text
        ocr_texts = []
        try:
            doc = fitz.open(stream=data, filetype="pdf")  # type: ignore
            for p in doc:
                pix = p.get_pixmap()  # type: ignore
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)  # type: ignore
                ocr_texts.append(pytesseract.image_to_string(img))  # type: ignore
            ot = "\n".join(ocr_texts)
            return ot if len(ot.strip()) > len(text.strip()) else text
        except Exception:
            return text
    except Exception as e:
        raise Exception(f"Error extracting PDF: {str(e)}")

def extract_text_from_docx(file):
    """Extract text from Word document"""
    try:
        doc = docx.Document(BytesIO(file.read()))
        paras = [paragraph.text for paragraph in doc.paragraphs]
        table_cells = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    val = cell.text.strip()
                    if val:
                        table_cells.append(val)
        text = "\n".join(paras + table_cells)
        return text
    except Exception as e:
        raise Exception(f"Error extracting DOCX: {str(e)}")

def extract_email(text):
    """Extract email from text"""
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b'
    emails = re.findall(email_pattern, text)
    return emails[0] if emails else None

def extract_phone(text):
    """Extract phone number with stricter heuristics.
    - Prefer lines labeled with phone/mobile/contact.
    - Exclude sequences inside emails/URLs.
    - Keep 10–15 digit numbers (with separators)."""
    # Remove email addresses to avoid picking numbers from them
    scrubbed = re.sub(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", " ", text)
    lines = scrubbed.splitlines()

    def normalize_digits(s: str) -> str:
        return re.sub(r"\D", "", s)

    # Candidate regex allowing separators, country code
    cand_re = re.compile(r"(?:\+\d{1,3}[-\s.]*)?(?:\(?\d{2,4}\)?[-\s.]*)?\d{3,4}[-\s.]*\d{4,6}")

    labeled = []
    unlabeled = []
    for ln in lines:
        if not ln.strip():
            continue
        for m in cand_re.finditer(ln):
            cand = m.group(0).strip()
            digits = normalize_digits(cand)
            if len(digits) < 10 or len(digits) > 15:
                continue
            if any(k in ln.lower() for k in ["phone", "mobile", "contact", "tel", "whatsapp"]):
                labeled.append(cand)
            else:
                unlabeled.append(cand)

    chosen = labeled[0] if labeled else (unlabeled[0] if unlabeled else None)
    return chosen

def extract_urls(text):
    """Extract LinkedIn/GitHub/Website URLs including bare domains without scheme."""
    urls = set()
    # Scheme URLs
    for m in re.findall(r'http[s]?://[^\s)]+', text):
        urls.add(m.strip(').,;'))
    # Bare domains (with optional path)
    for m in re.findall(r'\b(?:www\.)?[a-zA-Z0-9-]+\.(?:com|io|ai|dev|net|org)(?:/[^\s)]*)?', text):
        urls.add(m.strip(').,;'))

    linkedin = None
    github = None
    website = None

    # Normalize obvious LinkedIn/GitHub domains
    for url in list(urls):
        low = url.lower()
        # Normalize missing scheme
        normalized = url if low.startswith('http') else 'https://' + low
        if 'linkedin.com' in low:
            linkedin = normalized
        elif 'github.com' in low:
            github = normalized

    # Detect bare LinkedIn handles like 'in/username' not preceded by domain
    if not linkedin:
        m = re.search(r'\b(in/[\w\-]+)\b', text)
        if m:
            linkedin = 'https://www.linkedin.com/' + m.group(1)

    # Detect textual labels like 'LinkedIn: harshadthorat16' or 'GitHub: harshadpy'
    if not linkedin:
        m = re.search(r'(?i)linkedin\s*[:\-]?\s*([\w\-/]+)', text)
        if m:
            handle = m.group(1).strip().strip('.')
            if not handle.startswith('http'):
                if handle.startswith('in/') or handle.startswith('company/'):
                    linkedin = 'https://www.linkedin.com/' + handle
                else:
                    linkedin = 'https://www.linkedin.com/in/' + handle
            else:
                linkedin = handle

    if not github:
        # Handle without scheme, e.g., github.com/harshadpy captured above; else look for label
        m = re.search(r'(?i)github\s*[:\-]?\s*([\w\-/]+)', text)
        if m:
            handle = m.group(1).strip().strip('.')
            github = handle if handle.startswith('http') else f'https://github.com/{handle.lstrip("/")}'

    # Choose a general website that is not an email provider and not already LinkedIn/GitHub
    if urls and not website:
        for url in urls:
            low = url.lower()
            if any(dom in low for dom in ['linkedin.com', 'github.com', 'mailto:']):
                continue
            if any(dom in low for dom in ['gmail.com', 'yahoo.com', 'outlook.com', 'proton.me']):
                continue
            website = url if low.startswith('http') else 'https://' + low
            break

    return linkedin, github, website

def extract_name(text):
    """Extract name from text (first line typically)"""
    lines = text.strip().split('\n')
    for line in lines[:10]:
        line = line.strip()
        if not line:
            continue
        if any(k in line.lower() for k in ['email', 'phone', 'linkedin', 'github', 'www.', '@']):
            continue
        if any(k in line.lower() for k in ['resume', 'curriculum', 'vitae']):
            continue
        tokens = line.split()
        if 1 <= len(tokens) <= 5:
            return line
    return None

def extract_skills(text):
    """Extract skills from text.
    - Match from a Skills/Technical Skills section if present (multi-line).
    - Also scan against a known keyword list.
    - Split by common separators and clean tokens.
    - Dedupe and cap length."""
    skills_keywords = [
        'python', 'java', 'javascript', 'c++', 'c#', 'ruby', 'php', 'swift', 'kotlin',
        'react', 'angular', 'vue', 'node.js', 'express', 'django', 'flask', 'spring',
        'sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch',
        'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'git',
        'machine learning', 'deep learning', 'nlp', 'computer vision', 'data science',
        'tensorflow', 'pytorch', 'scikit-learn', 'pandas', 'numpy',
        'html', 'css', 'bootstrap', 'tailwind', 'sass',
        'agile', 'scrum', 'devops', 'ci/cd', 'microservices', 'rest api', 'graphql'
    ]

    def clean_token(tok: str) -> str:
        t = tok.strip().strip('•-*|,:;').strip()
        # Remove common noise
        t = re.sub(r"\b(skills?|technical|core|competencies)\b", "", t, flags=re.I).strip()
        # Normalize spacing
        t = re.sub(r"\s+", " ", t)
        return t

    found = []
    text_lower = text.lower()

    # Keyword scan
    for kw in skills_keywords:
        if kw in text_lower:
            found.append(kw)

    # Section scan (capture until next blank line or next section keyword)
    skills_section_pattern = r'(?is)(?:^|\n)\s*(?:skills?|technical skills?|core competencies)\s*[:\-]*\s*(.+?)(?=\n\s*\n|\n\s*(education|projects|experience)\b|$)'
    m = re.search(skills_section_pattern, text)
    if m:
        block = m.group(1)
        tokens = re.split(r'[\n,;|•]+', block)
        for tok in tokens:
            t = clean_token(tok)
            if not t:
                continue
            if len(t) > 40:
                continue
            # Keep up to 3 words tokens to avoid sentences, but allow common 2-word phrases
            if len(t.split()) <= 4:
                found.append(t)

    # Dedupe preserving order
    seen = set()
    uniq = []
    for s in found:
        key = s.lower()
        if key not in seen:
            seen.add(key)
            uniq.append(s)
    return [s if s.isupper() else s.title() for s in uniq[:50]]

def extract_education(text):
    """Extract education information"""
    education = []
    
    # Common degree keywords
    degree_keywords = [
        r'bachelor(?:\'s)?(?:\s+of)?(?:\s+science)?(?:\s+in)?',
        r'master(?:\'s)?(?:\s+of)?(?:\s+science)?(?:\s+in)?',
        r'phd', r'ph\.d\.', r'doctorate',
        r'associate(?:\'s)?(?:\s+of)?',
        r'b\.s\.', r'b\.a\.', r'm\.s\.', r'm\.a\.', r'mba'
    ]
    
    # Look for education section
    edu_section_pattern = r'(?i)(?:education|academic background)[:\s]*(.+?)(?=\n\n|\nexperience|\nwork experience|\nskills|\nprojects|$)'
    edu_match = re.search(edu_section_pattern, text, re.DOTALL)
    
    if edu_match:
        edu_text = edu_match.group(1)
        
        # Split by newlines and process
        lines = edu_text.split('\n')
        current_entry = ""
        
        for line in lines:
            line = line.strip()
            if line:
                current_entry += line + " "
                if any(re.search(keyword, current_entry, re.IGNORECASE) for keyword in degree_keywords) or re.search(r'\b\d{4}\b', current_entry):
                    education.append(current_entry.strip())
                    current_entry = ""
        
        if current_entry:
            education.append(current_entry.strip())
    
    return education if education else []

def extract_experience(text):
    """Extract work experience"""
    experience = []
    
    # Look for experience section
    exp_section_pattern = r'(?i)(?:experience|work experience|professional experience|employment)[:\s]*(.+?)(?=\n\n(?:education|skills|projects)|$)'
    exp_match = re.search(exp_section_pattern, text, re.DOTALL)
    
    if exp_match:
        exp_text = exp_match.group(1)
        
        # Split by job entries (typically separated by dates or company names)
        # This is a simplified extraction
        lines = [l for l in (exp_text.split('\n'))]
        current_job = {"title": "", "company": "", "dates": "", "responsibilities": []}
        
        for line in lines:
            line = line.strip()
            if not line:
                if current_job["title"]:
                    experience.append(current_job.copy())
                    current_job = {"title": "", "company": "", "dates": "", "responsibilities": []}
            elif re.search(r'(\b\d{1,2}\s*[A-Za-z]{3,9}\s*\d{4}\b|\b[A-Za-z]{3,9}\s*\d{4}\b|\b\d{4}\s*[-–]\s*(Present|\d{4})\b|\b\d{4}\b)', line):
                current_job["dates"] = line
            elif line.startswith(('•', '-', '*')):
                current_job["responsibilities"].append(line.lstrip('•-* '))
            else:
                if not current_job["title"] and not any(ch.isdigit() for ch in line):
                    current_job["title"] = line
                elif not current_job["company"]:
                    current_job["company"] = line
        
        if current_job["title"]:
            experience.append(current_job)
    
    return experience if experience else []

def extract_projects(text):
    projects = []
    section = re.search(r'(?is)(?:^|\n)\s*(projects?|personal projects?)\s*[:\-]*\s*(.+?)(?=\n\n|\n\s*(experience|education|skills)\b|$)', text)
    if section:
        body = section.group(2)
        blocks = re.split(r'\n\s*\n', body)
        for b in blocks:
            lines = [l.strip() for l in b.split('\n') if l.strip()]
            if not lines:
                continue
            name = lines[0]
            desc = ""
            tech = ""
            if len(lines) > 1:
                desc = lines[1]
            m = re.search(r'(?i)(tech|technologies|stack)\s*[:\-]\s*(.+)', b)
            if m:
                tech = m.group(2).strip()
            projects.append({"name": name, "description": desc, "technologies": tech})
    return projects

def parse_resume(text):
    """Main parser function to extract all resume information"""
    
    # Extract all components
    name = extract_name(text)
    email = extract_email(text)
    phone = extract_phone(text)
    linkedin, github, website = extract_urls(text)
    skills = extract_skills(text)
    education = extract_education(text)
    experience = extract_experience(text)
    projects = extract_projects(text)
    
    # Create structured resume data
    resume_data = {
        "personal_info": {
            "name": name or "Not found",
            "email": email or "Not found",
            "phone": phone or "",
            "linkedin": linkedin or "",
            "github": github or "",
            "website": website or "",
            "location": "",
            "summary": ""
        },
        "skills": skills,
        "education": education,
        "experience": experience,
        "projects": projects
    }

    try:
        resume_data = _enrich_with_ner_and_heuristics(text, resume_data)
    except Exception:
        pass
    return resume_data


def _enrich_with_ner_and_heuristics(text, data):
    loc = data.get("personal_info", {}).get("location") or ""
    nm = data.get("personal_info", {}).get("name") or ""
    nlp = None
    if spacy:
        try:
            nlp = spacy.load("en_core_web_sm")  # type: ignore
        except Exception:
            try:
                nlp = spacy.blank("en")  # type: ignore
            except Exception:
                nlp = None
    if nlp:
        doc = nlp(text[:2000])
        if not nm:
            for ent in doc.ents:
                if ent.label_ == "PERSON" and 2 <= len(ent.text.split()) <= 5:
                    nm = ent.text
                    break
        if not loc:
            for ent in doc.ents:
                if ent.label_ in ("GPE", "LOC"):
                    loc = ent.text
                    break
    if not loc:
        m = re.search(r"(?i)\b(location|based in|residing in|address)\b\s*[:\-]?\s*(.+)", text)
        if m:
            loc = m.group(2).split("\n")[0].strip().strip(",")
    if not loc:
        m = re.search(r"\b([A-Za-z .]+,\s*[A-Za-z .]+)\b", text)
        if m and len(m.group(1)) <= 40:
            loc = m.group(1).strip()
    if nm:
        data["personal_info"]["name"] = nm
    if loc:
        data["personal_info"]["location"] = loc
    return data